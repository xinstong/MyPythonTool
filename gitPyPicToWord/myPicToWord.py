# -*- coding: UTF-8 -*-
import os
import time
import glob
import re
from threading import Thread
from wx.lib.pubsub import pub
import win32com.client as win32
from win32com.client import constants
from docx import Document
from docx.shared import Inches
from PIL import Image

try:
    import wx
except ImportError:
    raise ImportError("we need wxPython.")

class wxPicWordTool(wx.Frame):
    def __init__(self, parent, id, title):
        wx.Frame.__init__(self, parent, id, title)
        self.par = parent
        self.initialize()

    def initialize(self):
        panel = wx.Panel(self, size=wx.Size(800, 600))
        sizer = wx.GridBagSizer(0, 0)

        self.tcPath = wx.TextCtrl(panel, size=wx.Size(250, 20))
        self.Bind(wx.EVT_TEXT, self.onTcPathTextChged, self.tcPath)
        sizer.Add(self.tcPath, pos=(0, 0), span=(1, 1), flag=wx.ALL|wx.EXPAND, border=5)

        self.btBrowse=wx.Button(panel, label=u"浏览")
        self.Bind(wx.EVT_BUTTON, self.OnButtonBrowse, self.btBrowse)
        sizer.Add(self.btBrowse, pos=(0, 1), flag=wx.ALIGN_TOP|wx.ALIGN_RIGHT, border=5)

        self.gauge = wx.Gauge(panel, range=20, size=(250, 20), style=wx.GA_HORIZONTAL)
        sizer.Add(self.gauge, pos=(1, 0), flag=wx.EXPAND|wx.ALL, border=5)
        pub.subscribe(self.updateProgress, "updateDegree")

        self.btHandle=wx.Button(panel, label=u"开始处理")
        self.btHandle.Enable(False)
        self.Bind(wx.EVT_BUTTON, self.OnButtonHandle, self.btHandle)
        sizer.Add(self.btHandle, pos=(1, 1), flag=wx.ALIGN_TOP|wx.ALIGN_RIGHT, border=5)
        sizer.SetFlexibleDirection(wx.BOTH);
        sizer.SetNonFlexibleGrowMode(wx.FLEX_GROWMODE_ALL)

        self.tcResult=wx.TextCtrl(panel, style=wx.TE_MULTILINE|wx.TE_READONLY, size=wx.Size(250, 20))
        sizer.Add(self.tcResult, pos=(2, 0), span=(1, 1), flag=wx.EXPAND|wx.ALL, border = 5)
        pub.subscribe(self.reportResult, "reportResult")

        sizer.AddGrowableCol(0)
        sizer.AddGrowableRow(2)

        panel.SetSizerAndFit(sizer)
        self.SetSizeHints(700,400,800,700)
        self.Centre()
        self.CreateStatusBar()
        self.SetStatusText(u"欢迎使用转换工具!")
        self.Show(True)

    def OnButtonBrowse(self, event):
        dirDlg=wx.DirDialog(self, message=u"选择待处理文件夹", style=wx.DD_DEFAULT_STYLE | wx.DD_DIR_MUST_EXIST)
        if wx.ID_OK == dirDlg.ShowModal():
            self.tcPath.SetLabelText(dirDlg.GetPath())

    def onTcPathTextChged(self, event):
        self.btHandle.Enable(os.path.exists(self.tcPath.Value))

    def OnButtonHandle(self, event):
        self.btHandle.Enable(False)
        self.btBrowse.Enable(False)
        self.tcPath.SetEditable(False)
        str_text = "----------------"
        str_text += time.strftime("%a, %d %b %Y %H:%M:%S +0000", time.gmtime())
        str_text += "---------------\n"
        self.tcResult.AppendText(str_text)
        print("will start thread" + repr(self.tcPath.Value))
        MyWorkThred(self.tcPath.Value)

    def updateProgress(self, allTask, curStep):
        self.gauge.SetValue(curStep)
        self.gauge.SetRange(allTask)
        if curStep == allTask:
            self.btHandle.Enable(True)
            self.btBrowse.Enable(True)
            self.tcPath.SetEditable(True)

    def reportResult(self, allTask, curStep, docName, picName, result):
        print("allTask {0:}, curStep {1:}".format(allTask, curStep))
        str_res=""
        str_res += "["
        str_res += repr(curStep)
        str_res += "/"
        str_res += repr(allTask)
        str_res += "]:"
        if 0 == result:
            str_res += "Ok"
        else:
            str_res += "NG"
        str_res += " ["
        str_res += picName
        str_res += "] to ["
        str_res += docName
        str_res += "].\n"

        self.tcResult.AppendText(str_res)


class MyWorkThred(Thread):
    taskCount=0
    curStep=0
    rootPath=""
    wordPaths=list()
    picPaths=list()
    def __init__(self, rootPath):
        self.rootPath=rootPath
        Thread.__init__(self)
        self.start()

    def gatherAll(self):
        docFiles = glob.glob(self.rootPath + '/**/*.doc', recursive=True)
        docxFiles = glob.glob(self.rootPath + '/**/*.docx', recursive=True)
        pngFiles = glob.glob(self.rootPath + '/**/*.png', recursive=True)
        newDocxFiles=list()
        for docPath in docFiles:
            word = win32.gencache.EnsureDispatch('Word.Application')
            docWord = word.Documents.Open(docPath)
            docWord.Activate()

            new_file_abs = os.path.abspath(docPath)
            new_file_abs = re.sub(r'\.\w+$', '.docx', new_file_abs)

            word.ActiveDocument.SaveAs(
                new_file_abs, FileFormat=constants.wdFormatXMLDocument
            )
            docWord.Close(False)
            newDocxFiles.append(new_file_abs)
            os.remove(docPath)

        jpgFiles = glob.glob(self.rootPath + '/**/*.jpg', recursive=True)
        self.wordPaths.clear()
        for doc in newDocxFiles:
            self.wordPaths.append(doc)
        for docx in docxFiles:
            self.wordPaths.append(docx)
        self.taskCount = len(self.wordPaths)

        self.picPaths.clear()
        for png in pngFiles:
            self.picPaths.append(png)
        for jpg in jpgFiles:
            self.picPaths.append(jpg)
        self.curStep = 0

    def run(self):
        print("run")
        self.gatherAll()
        for wordPath in self.wordPaths:
            self.curStep += 1
            str_word_path = r''
            str_word_path = wordPath
            short_word_path = str_word_path[len(self.rootPath) : len(str_word_path)]
            str_word_keys = str_word_path.split('.')
            if len(str_word_keys) == 0:
                print("ERROR: path[" + str_word_path + "]can't split")
                continue
            str_pic_path = ""
            bFinded = False
            for picPath in self.picPaths:
                str_pic_path = picPath
                if str_word_keys[0] in str_pic_path:
                    bFinded = True
                    break
            if not bFinded:
                print("ERROR: path[" + str_word_path + "]can't get pic")
                continue
            document = Document(str_word_path)

            p = document.add_paragraph()
            r = p.add_run()
            image = Image.open(str_pic_path)
            image.rotate(90)
            myOldWidth, myOldHeight = image.size
            print("myOldWidth{0:}, myOldHeight{1:}".format(myOldWidth, myOldHeight))
            myNewWidth = 100
            myNewHeight = myNewWidth / myOldWidth * myOldHeight
            print("myNewWidth{0:}, myNewHeight{1:}".format(myNewWidth, myNewHeight))
            myNewSize = (myNewWidth, myNewHeight)
            #image.thumbnail(myNewSize, Image.ANTIALIAS)
            image.save(str_pic_path)
            r.add_picture(str_pic_path, width=Inches(4.7), height=Inches(5.7))
            document.save(str_word_path)

            short_word_path = str_word_path[len(self.rootPath) : len(str_word_path)]
            short_pic_path = str_pic_path[len(self.rootPath) : len(str_pic_path)]

            time.sleep(0.01)
            pub.sendMessage('updateDegree', allTask=self.taskCount, curStep=self.curStep)
            pub.sendMessage('reportResult', allTask=self.taskCount, curStep=self.curStep,
                            docName=short_word_path, picName=short_pic_path, result=0)

if __name__=="__main__":
    app = wx.App()
    frame = wxPicWordTool(None, -1, u"插入图片到word")
    app.MainLoop()
